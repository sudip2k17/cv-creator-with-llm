"""
cv_creator_capstone_extended.py

Capstone Project: CV Creation using LLMs
Now integrates:
1. Gemma 3 1B via Ollama (core tailoring)
2. LlamaIndex (resume/job parsing)
3. ResumeLM (optional ATS-optimized generation/validation)
4. LangChain (orchestration of pipeline)

Author: Sudip Sengupta
"""

import os
import json
import re
import requests
import pdfplumber
import docx
from pathlib import Path

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
    print("✅ Environment variables loaded from .env file")
except ImportError:
    print("⚠️ python-dotenv not available. Install with: pip install python-dotenv")

# --- OpenAI API Configuration (DISABLED) ---
# OpenAI quota exceeded - using local alternatives instead
OPENAI_API_KEY = None
print("⚠️ OpenAI API disabled due to quota limits - using local alternatives")

# --- LLM & Orchestration ---
try:
    from langchain_community.llms import Ollama
    from langchain.prompts import PromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except ImportError:
    print("⚠️ LangChain not available")
    Ollama = None
    PromptTemplate = None
    LLMChain = None
    LANGCHAIN_AVAILABLE = False

# --- LlamaIndex with Local Embeddings ---
try:
    from llama_index.core import Document, VectorStoreIndex, Settings
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    
    # Configure LlamaIndex to use local HuggingFace embeddings
    Settings.embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
    # Use Ollama for LLM instead of OpenAI
    Settings.llm = None  # We'll handle LLM calls separately
    
    print("✅ LlamaIndex configured with local HuggingFace embeddings")
    LLAMAINDEX_AVAILABLE = True
except ImportError as e:
    print("⚠️ LlamaIndex not available. Install with: pip install llama-index llama-index-embeddings-huggingface sentence-transformers")
    print(f"   Error: {e}")
    Document = None
    VectorStoreIndex = None
    Settings = None
    LLAMAINDEX_AVAILABLE = False

# --- ResumeLM (optional - assume cloned repo or pip install) ---
# NOTE: ResumeLM is not on PyPI by default. Clone from GitHub and ensure it's in PYTHONPATH.
try:
    from resumelm import ResumeBuilder, ATSScorer
    HAS_RESUMELM = True
except ImportError:
    print("⚠️ ResumeLM not available (optional). Skipping advanced ATS features.")
    ResumeBuilder = None
    ATSScorer = None
    HAS_RESUMELM = False

# Configuration
OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "gemma3:1b")


# =============================
# Utility Functions
# =============================

def find_working_ollama_url():
    """Find a working Ollama URL from common options"""
    test_urls = [
        "http://localhost:11434",
        "http://127.0.0.1:11434", 
        "http://0.0.0.0:11434"
    ]
    
    for url in test_urls:
        try:
            response = requests.get(f"{url}/api/version", timeout=2)
            if response.status_code == 200:
                return url
        except:
            continue
    
    return OLLAMA_URL  # Return default as fallback

def call_ollama_direct(prompt: str, model: str = OLLAMA_MODEL, temperature: float = 0.0, max_tokens: int = 1024):
    """Direct Ollama API call with better error handling"""
    # Try to find working URL dynamically
    working_url = find_working_ollama_url()
    url = f"{working_url}/api/generate"
    
    payload = {
        "model": model, 
        "prompt": prompt, 
        "stream": False,  # Important: disable streaming for simpler handling
        "options": {
            "temperature": temperature, 
            "num_predict": max_tokens
        }
    }
    
    try:
        resp = requests.post(url, json=payload, timeout=120)
        resp.raise_for_status()
        data = resp.json()
        
        # Handle Ollama's response format
        if isinstance(data, dict) and 'response' in data:
            return data['response']
        elif isinstance(data, dict):
            # Try other possible response fields
            if 'results' in data and isinstance(data['results'], list):
                return data['results'][0].get('content', '')
            if 'text' in data:
                return data['text']
            if 'result' in data:
                return data['result']
        
        return str(data)
        
    except requests.exceptions.ConnectionError:
        return "Error: Cannot connect to Ollama. Please make sure Ollama is running and accessible."
    except requests.exceptions.Timeout:
        return "Error: Request timed out. The model might be taking too long to respond."
    except Exception as e:
        return f"Error calling Ollama: {e}"

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract raw text from PDF."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(docx_path: str) -> str:
    """Extract raw text from DOCX."""
    doc = docx.Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs])


# =============================
# Model Components
# =============================

# 1. Gemma 3 1B via Ollama - Use direct API calls instead of LangChain
USE_LANGCHAIN = False
ollama_llm = None

if LANGCHAIN_AVAILABLE:
    try:
        ollama_llm = Ollama(model=OLLAMA_MODEL)
        print("✅ LangChain Ollama initialized")
    except Exception as e:
        print(f"⚠️ Could not initialize LangChain Ollama: {e}")
        print("   Using direct API calls instead")
        USE_LANGCHAIN = False
        ollama_llm = None
else:
    print("🔧 LangChain disabled - using direct Ollama API only")

# 2. LlamaIndex-based parser with local fallback
def parse_with_llama_index(text: str, query: str) -> str:
    """Parse text using LlamaIndex or fallback to simple processing"""
    
    # Fallback 1: If LlamaIndex is not available, use simple extraction
    if not LLAMAINDEX_AVAILABLE or Document is None or VectorStoreIndex is None:
        print("⚠️ LlamaIndex not available, using simple text extraction")
        return simple_text_extraction(text, query)
    
    # Try LlamaIndex with local embeddings
    try:
        print(f"🔍 Processing with LlamaIndex: {query[:50]}...")
        doc = Document(text=text)
        index = VectorStoreIndex.from_documents([doc])
        query_engine = index.as_query_engine()
        response = query_engine.query(query)
        return getattr(response, 'response', str(response))
    except Exception as e:
        print(f"⚠️ LlamaIndex parsing failed: {e}")
        print("   Falling back to simple text extraction...")
        return simple_text_extraction(text, query)


def simple_text_extraction(text: str, query: str) -> str:
    """Simple text extraction when LlamaIndex is not available"""
    
    query_lower = query.lower()
    
    if "name" in query_lower and "contact" in query_lower:
        # Extract first few lines for contact info
        lines = text.split('\n')[:5]
        return '\n'.join([line.strip() for line in lines if line.strip()])
    elif "skills" in query_lower:
        # Look for skills
        skill_keywords = ['python', 'java', 'javascript', 'react', 'sql', 'html', 'css', 'git']
        found_skills = []
        for line in text.split('\n'):
            for skill in skill_keywords:
                if skill.lower() in line.lower():
                    found_skills.append(skill.title())
        return f'{{"skills": {list(set(found_skills))}}}'
    elif "experience" in query_lower or "work" in query_lower:
        # Look for experience-related content
        exp_keywords = ['experience', 'work', 'employment', 'position', 'role', 'company']
        relevant_lines = []
        for line in text.split('\n'):
            if any(keyword in line.lower() for keyword in exp_keywords):
                relevant_lines.append(line.strip())
        return '\n'.join(relevant_lines[:15])
    elif "education" in query_lower:
        # Look for education-related content
        edu_keywords = ['education', 'degree', 'university', 'college', 'school', 'bachelor', 'master']
        relevant_lines = []
        for line in text.split('\n'):
            if any(keyword in line.lower() for keyword in edu_keywords):
                relevant_lines.append(line.strip())
        return '\n'.join(relevant_lines[:10])
    elif "responsibilities" in query_lower or "requirements" in query_lower:
        # Extract job-related content
        job_keywords = ['responsibilities', 'requirements', 'qualifications', 'duties', 'skills required']
        relevant_lines = []
        for line in text.split('\n'):
            if any(keyword in line.lower() for keyword in job_keywords):
                relevant_lines.append(line.strip())
        return '\n'.join(relevant_lines[:20])
    else:
        # Default: return first portion of text
        return text[:500]


# 3. Resume Tailoring with LangChain or Direct API
if LANGCHAIN_AVAILABLE and PromptTemplate is not None and LLMChain is not None:
    resume_prompt = PromptTemplate(
        input_variables=["resume_json", "job_json"],
        template="""
You are an expert CV writer. 
Tailor the following resume JSON to match the given job JSON.
Ensure ATS-friendliness and keep it concise, keyword-rich, and professional.

Resume: {resume_json}
Job: {job_json}

Return a polished resume in JSON format with sections: name, contact, summary, skills, experience, education, projects.
"""
    )

    # Initialize the chain only if Ollama is available
    if ollama_llm is not None and USE_LANGCHAIN:
        try:
            resume_chain = LLMChain(llm=ollama_llm, prompt=resume_prompt)
            print("✅ LangChain resume chain initialized")
        except Exception as e:
            print(f"⚠️ Could not initialize LangChain: {e}")
            resume_chain = None
    else:
        resume_chain = None
else:
    resume_prompt = None
    resume_chain = None

# Simple safe JSON extractor
def safe_json_from_text(text: str):
    """Extract JSON from text response"""
    if not isinstance(text, str):
        return text
    m = re.search(r"(\{[\s\S]*\})", text)
    if m:
        blob = m.group(1)
        try:
            return json.loads(blob)
        except Exception:
            try:
                return json.loads(blob.replace("'", '"'))
            except Exception:
                return blob
    # fallback
    return text

# ATS keyword score
def ats_keyword_score(resume_text: str, keywords: list):
    """Calculate ATS keyword matching score"""
    if not keywords:
        return {"matched": [], "score": 0.0}
    rl = resume_text.lower()
    found = [k for k in keywords if k.lower() in rl]
    score = len(found)/len(keywords)
    return {"matched": found, "score": round(score,3)}


# 4. Optional ResumeLM ATS generation
def generate_with_resumelm(parsed_resume, job_desc):
    """Generate resume using ResumeLM if available"""
    if not HAS_RESUMELM or ResumeBuilder is None:
        return None
    try:
        builder = ResumeBuilder()
        ats_resume = builder.build(parsed_resume, job_desc)
        return ats_resume
    except Exception as e:
        print(f"⚠️ ResumeLM generation failed: {e}")
        return None


def score_with_resumelm(resume_text, job_desc):
    """Score resume using ResumeLM if available"""
    if not HAS_RESUMELM or ATSScorer is None:
        return None
    try:
        scorer = ATSScorer()
        return scorer.score(resume_text, job_desc)
    except Exception as e:
        print(f"⚠️ ResumeLM scoring failed: {e}")
        return None


# =============================
# Main Pipeline
# =============================

def build_cv(resume_file: str, job_file: str, output_path: str = "final_resume.json"):
    """Process a single resume against a job description"""
    print(f"🚀 Building CV for {resume_file}...")
    
    # --- Step 1: Load Resume & Job ---
    print("📄 Loading resume and job description...")
    if resume_file.endswith(".pdf"):
        resume_text = extract_text_from_pdf(resume_file)
    else:
        resume_text = extract_text_from_docx(resume_file)

    job_text = Path(job_file).read_text(encoding='utf-8')

    # --- Step 2: Parse with LlamaIndex (or fallback) ---
    print("🔍 Parsing resume and job with LlamaIndex...")
    parsed_resume = parse_with_llama_index(
        resume_text, "Extract name, contact, skills, education, work experience, and projects in JSON."
    )
    parsed_job = parse_with_llama_index(
        job_text, "Extract required skills, role responsibilities, and keywords in JSON."
    )

    # --- Step 3: Tailor Resume with Direct Ollama API (preferred) or LangChain ---
    print("✨ Tailoring resume with Ollama...")
    
    # Try direct API first (more reliable)
    prompt = f"""You are an expert CV writer. Create a tailored resume based on the following:

PARSED RESUME DATA:
{parsed_resume}

JOB REQUIREMENTS:
{parsed_job}

Please create a professional, ATS-friendly resume in JSON format that:
1. Highlights skills matching the job requirements
2. Emphasizes relevant experience
3. Uses keywords from the job description
4. Maintains a clean, professional structure
5. Includes sections: name, contact, summary, skills, experience, education, projects

Return only the JSON resume without any additional commentary."""

    tailored_resume_raw = call_ollama_direct(prompt, model=OLLAMA_MODEL, temperature=0.2, max_tokens=2000)
    
    # Try to extract JSON from response
    tailored_resume_json = safe_json_from_text(tailored_resume_raw)
    
    # Fallback to LangChain if direct API fails and LangChain is available
    if isinstance(tailored_resume_json, str) and "Error:" in tailored_resume_json and resume_chain is not None:
        print("⚠️ Direct API failed, trying LangChain...")
        try:
            tailored_resume_json = resume_chain.run(
                {"resume_json": parsed_resume, "job_json": parsed_job}
            )
            tailored_resume_json = safe_json_from_text(tailored_resume_json)
        except Exception as e:
            print(f"⚠️ LangChain processing failed: {e}")
            tailored_resume_json = {
                "error": "Both direct API and LangChain failed",
                "resume": parsed_resume, 
                "job": parsed_job,
                "raw_response": tailored_resume_raw
            }
    elif isinstance(tailored_resume_json, str) and "Error:" in tailored_resume_json:
        print("⚠️ Resume tailoring failed, using basic combination")
        tailored_resume_json = {
            "error": "Resume tailoring failed", 
            "resume": parsed_resume, 
            "job": parsed_job,
            "raw_response": tailored_resume_raw
        }

    # --- Step 4: Optionally enhance with ResumeLM ---
    print("🎯 Enhancing with ResumeLM (if available)...")
    ats_resume = tailored_resume_json
    ats_score = None
    
    if HAS_RESUMELM:
        try:
            resumelm_result = generate_with_resumelm(tailored_resume_json, parsed_job)
            if resumelm_result:
                ats_resume = resumelm_result
                ats_score = score_with_resumelm(json.dumps(ats_resume), job_text)
                print("📊 ATS Score (ResumeLM):", ats_score)
        except Exception as e:
            print(f"⚠️ ResumeLM processing failed: {e}")
    
    # --- Step 5: Calculate basic ATS score ---
    print("📊 Calculating ATS keyword score...")
    keywords = []
    parsed_job_data = safe_json_from_text(parsed_job)
    if isinstance(parsed_job_data, dict):
        keywords = parsed_job_data.get('keywords') or parsed_job_data.get('requirements') or []
        if isinstance(keywords, str):
            keywords = [k.strip() for k in re.split(r'[,;]|\n', keywords) if k.strip()]
    
    basic_score = ats_keyword_score(json.dumps(ats_resume), keywords)
    print(f"📊 Basic ATS Score: {basic_score}")

    # --- Step 6: Prepare final output ---
    final_output = {
        "tailored_resume": ats_resume,
        "parsing_results": {
            "parsed_resume": parsed_resume,
            "parsed_job": parsed_job
        },
        "ats_scores": {
            "basic_score": basic_score,
            "resumelm_score": ats_score
        },
        "metadata": {
            "ollama_model": OLLAMA_MODEL,
            "langchain_available": LANGCHAIN_AVAILABLE,
            "llamaindex_available": LLAMAINDEX_AVAILABLE,
            "resumelm_available": HAS_RESUMELM,
            "source_resume": resume_file
        }
    }

    # --- Step 7: Save Output ---
    print(f"💾 Saving output to {output_path}...")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(json.dumps(final_output, indent=2))

    print(f"✅ Final tailored resume saved to {output_path}")
    return final_output


def build_cv_batch(resume_files: list, job_file: str, output_dir: str = "tailored_resumes"):
    """Process multiple resumes against a single job description"""
    import time
    from datetime import datetime
    
    print(f"🚀 Starting batch CV processing for {len(resume_files)} resumes...")
    print(f"📁 Output directory: {output_dir}")
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Process resumes in batches of 10
    batch_size = 10
    total_resumes = len(resume_files)
    successful_outputs = []
    failed_outputs = []
    
    for batch_start in range(0, total_resumes, batch_size):
        batch_end = min(batch_start + batch_size, total_resumes)
        batch_resumes = resume_files[batch_start:batch_end]
        
        print(f"\n📦 Processing batch {batch_start//batch_size + 1} ({batch_start + 1}-{batch_end} of {total_resumes})...")
        
        for i, resume_file in enumerate(batch_resumes):
            resume_index = batch_start + i + 1
            
            try:
                print(f"\n--- Processing Resume {resume_index}/{total_resumes}: {resume_file} ---")
                
                # Generate output filename
                resume_name = Path(resume_file).stem
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"tailored_{resume_name}_{timestamp}.json"
                output_path = os.path.join(output_dir, output_filename)
                
                # Process the resume
                result = build_cv(resume_file, job_file, output_path)
                successful_outputs.append({
                    "resume": resume_file,
                    "output": output_path,
                    "basic_score": result.get("ats_scores", {}).get("basic_score", {}),
                    "resumelm_score": result.get("ats_scores", {}).get("resumelm_score")
                })
                
                print(f"✅ Resume {resume_index} processed successfully")
                
                # Small delay between processing to avoid overwhelming the system
                time.sleep(1)
                
            except Exception as e:
                print(f"❌ Error processing resume {resume_index} ({resume_file}): {e}")
                failed_outputs.append({
                    "resume": resume_file,
                    "error": str(e)
                })
        
        # Longer pause between batches
        if batch_end < total_resumes:
            print(f"⏸️  Batch complete. Pausing 3 seconds before next batch...")
            time.sleep(3)
    
    # Generate summary report
    print(f"\n📊 Batch Processing Summary:")
    print(f"   Total resumes: {total_resumes}")
    print(f"   Successful: {len(successful_outputs)}")
    print(f"   Failed: {len(failed_outputs)}")
    
    # Save summary report
    summary_report = {
        "batch_summary": {
            "total_resumes": total_resumes,
            "successful": len(successful_outputs),
            "failed": len(failed_outputs),
            "processed_at": datetime.now().isoformat(),
            "job_file": job_file,
            "output_directory": output_dir
        },
        "successful_outputs": successful_outputs,
        "failed_outputs": failed_outputs
    }
    
    summary_path = os.path.join(output_dir, f"batch_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(json.dumps(summary_report, indent=2))
    
    print(f"📋 Summary report saved to: {summary_path}")
    
    if successful_outputs:
        print(f"\n🎯 Top performing resumes by ATS score:")
        sorted_results = sorted(successful_outputs, 
                              key=lambda x: x.get("basic_score", {}).get("score", 0), 
                              reverse=True)
        for i, result in enumerate(sorted_results[:5]):
            score = result.get("basic_score", {}).get("score", 0)
            print(f"   {i+1}. {Path(result['resume']).name} - Score: {score:.3f}")
    
    return summary_report


def find_resume_files(directory: str, extensions: list = ['.pdf', '.docx']) -> list:
    """Find all resume files in a directory"""
    resume_files = []
    directory_path = Path(directory)
    
    if not directory_path.exists():
        print(f"❌ Directory not found: {directory}")
        return []
    
    for ext in extensions:
        resume_files.extend(directory_path.glob(f"*{ext}"))
        resume_files.extend(directory_path.glob(f"**/*{ext}"))  # Include subdirectories
    
    # Convert to strings and remove duplicates
    resume_files = list(set([str(f) for f in resume_files]))
    resume_files.sort()  # Sort for consistent processing order
    
    print(f"📁 Found {len(resume_files)} resume files in {directory}")
    return resume_files


# =============================
# CLI Usage
# =============================
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="CV Creator Capstone with Gemma, LlamaIndex, ResumeLM, LangChain - Single or Batch Processing")
    
    # Mode selection
    parser.add_argument("--mode", choices=["single", "batch"], default="single", 
                       help="Processing mode: 'single' for one resume, 'batch' for multiple resumes")
    
    # Single mode arguments
    parser.add_argument("--resume", help="Path to input resume (PDF/DOCX) - for single mode")
    parser.add_argument("--job", required=True, help="Path to job description (TXT)")
    parser.add_argument("--out", default="final_resume.json", help="Output JSON file - for single mode")
    
    # Batch mode arguments
    parser.add_argument("--resume-dir", help="Directory containing resume files (PDF/DOCX) - for batch mode")
    parser.add_argument("--output-dir", default="tailored_resumes", help="Output directory for batch mode")
    parser.add_argument("--extensions", nargs="+", default=[".pdf", ".docx"], 
                       help="File extensions to search for (default: .pdf .docx)")
    
    args = parser.parse_args()
    
    if args.mode == "single":
        # Single resume processing
        if not args.resume:
            print("❌ Error: --resume is required for single mode")
            parser.print_help()
            exit(1)
        
        print("🔧 Single Resume Processing Mode")
        build_cv(args.resume, args.job, args.out)
        
    elif args.mode == "batch":
        # Batch resume processing
        if not args.resume_dir:
            print("❌ Error: --resume-dir is required for batch mode")
            parser.print_help()
            exit(1)
        
        print("🔧 Batch Resume Processing Mode")
        
        # Find all resume files in directory
        resume_files = find_resume_files(args.resume_dir, args.extensions)
        
        if not resume_files:
            print("❌ No resume files found in the specified directory")
            exit(1)
        
        print(f"📁 Found {len(resume_files)} resume files")
        
        # Confirm before processing
        user_input = input(f"Do you want to process {len(resume_files)} resumes? (y/N): ")
        if user_input.lower() not in ['y', 'yes']:
            print("❌ Batch processing cancelled")
            exit(0)
        
        # Process batch
        summary = build_cv_batch(resume_files, args.job, args.output_dir)
        
        print(f"\n🎉 Batch processing complete!")
        print(f"   Processed: {summary['batch_summary']['successful']}/{summary['batch_summary']['total_resumes']} resumes")
        print(f"   Output directory: {args.output_dir}")
    
    print("\n✅ Processing complete!")
